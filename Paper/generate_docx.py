import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def create_element(name):
    return OxmlElement(name)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_borders(cell, color="CCCCCC", sz="4", val="single"):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), val)
        border.set(qn('w:sz'), sz)
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)

doc = docx.Document()

# Set Page Size to A4 and margins to 2.54 cm (1 inch)
for section in doc.sections:
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

# Configure normal style
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
font.color.rgb = RGBColor(0, 0, 0)

# Header Jurnal
p_hdr = doc.add_paragraph()
p_hdr.alignment = WD_ALIGN_PARAGRAPH.LEFT
p_hdr_run = p_hdr.add_run("Jurnal Pendidikan: Teori, Penelitian, dan Pengembangan\nVolume: 11 Number: 2 Month-Year: Juni 2026 Page: 120-130\nAvailable online http://journal.um.ac.id/index.php/jptpp/\nEISSN: 2502-471X | DOAJ-SHERPA/RoMEO-Google Scholar-IPI")
p_hdr_run.font.size = Pt(9)
p_hdr_run.font.italic = True

# Spasi sebelum judul
doc.add_paragraph()

# Judul
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_title = p_title.add_run("PENGEMBANGAN LEARNING MANAGEMENT SYSTEM BIMBINGAN BAHASA ARAB BERBASIS WEB MENGGUNAKAN FRAMEWORK LARAVEL PADA BAIT TAHFIZ AL-QURAN RIDHALLAH")
run_title.font.size = Pt(14)
run_title.font.bold = True

# Penulis
p_author = doc.add_paragraph()
p_author.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_author = p_author.add_run("Hanif Maulana Aldyza¹*, Nazaruddin Ahmad²")
run_author.font.bold = True
run_author.font.size = Pt(12)

# Afiliasi
p_aff = doc.add_paragraph()
p_aff.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_aff = p_aff.add_run("¹Program Studi Pendidikan Teknologi Informasi, Universitas Islam Negeri Ar-Raniry Banda Aceh, Indonesia\n²Fakultas Tarbiyah dan Keguruan, Universitas Islam Negeri Ar-Raniry Banda Aceh, Indonesia\nEmail: hanif.aldyza@ar-raniry.ac.id*, nazaruddin.ahmad@ar-raniry.ac.id")
run_aff.font.size = Pt(10)
run_aff.font.italic = True

doc.add_paragraph()

# Table for Article Info and Abstract
table = doc.add_table(rows=1, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = False

col_widths = [Inches(2.2), Inches(4.0)]
for i, col in enumerate(table.columns):
    col.width = col_widths[i]

cell_info = table.cell(0, 0)
cell_abs = table.cell(0, 1)

cell_info.width = Inches(2.2)
cell_abs.width = Inches(4.0)

set_cell_margins(cell_info, top=140, bottom=140, left=140, right=140)
set_cell_margins(cell_abs, top=140, bottom=140, left=140, right=140)
set_cell_borders(cell_info, color="999999", sz="4")
set_cell_borders(cell_abs, color="999999", sz="4")

# Column 1: Article Info
p_info_title = cell_info.paragraphs[0]
p_info_title.paragraph_format.space_after = Pt(6)
run_info_title = p_info_title.add_run("ARTICLE INFO")
run_info_title.font.bold = True
run_info_title.font.size = Pt(10)

p_info_body = cell_info.add_paragraph()
p_info_body.paragraph_format.space_after = Pt(6)
p_info_body.paragraph_format.line_spacing = 1.15
run_info_body = p_info_body.add_run(
    "Article history:\n"
    "Accepted: 10 Mei 2026\n"
    "Approved: 5 Juni 2026\n\n"
    "Keywords:\n"
    "Learning Management System,\n"
    "Laravel, Vue.js,\n"
    "Bimbingan Bahasa Arab,\n"
    "Waterfall\n\n"
    "Author Correspondence:\n"
    "Hanif Maulana Aldyza\n"
    "UIN Ar-Raniry Banda Aceh\n"
    "Email: hanif.aldyza@ar-raniry.ac.id"
)
run_info_body.font.size = Pt(9)

# Column 2: Abstract
p_abs_title = cell_abs.paragraphs[0]
p_abs_title.paragraph_format.space_after = Pt(6)
p_abs_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_abs_title = p_abs_title.add_run("ABSTRACT")
run_abs_title.font.bold = True
run_abs_title.font.size = Pt(10)

p_abs_body = cell_abs.add_paragraph()
p_abs_body.paragraph_format.line_spacing = 1.15
run_abs_body = p_abs_body.add_run(
    "Penelitian ini bertujuan untuk merancang dan mengembangkan Learning Management System (LMS) berbasis web untuk bimbingan bahasa Arab di Bait Tahfizh Al-Qur'an Ridhallah (BTQR) guna mengatasi kendala geografis peserta luar daerah yang tidak dapat mengikuti pembelajaran tatap muka. Pengembangan dilakukan dengan menggunakan metode Waterfall yang meliputi tahapan analisis kebutuhan, perancangan sistem, implementasi kode, pengujian, peluncuran, dan pemeliharaan. LMS dibangun secara kolaboratif menggunakan Laravel pada sisi backend dan Vue.js pada sisi frontend dengan mengadopsi rancangan antarmuka teruji. Pengujian fungsionalitas sistem dilakukan secara menyeluruh menggunakan metode black-box testing. Hasil penelitian menunjukkan bahwa seluruh modul utama, termasuk manajemen akun, pengelolaan modul materi, kuis, dan pelaksanaan try out, berfungsi dengan sangat baik sesuai spesifikasi kebutuhan. Dengan demikian, platform LMS ini berhasil memperluas aksesibilitas dan meningkatkan efektivitas bimbingan bahasa Arab di BTQR secara jarak jauh."
)
run_abs_body.font.size = Pt(10)
run_abs_body.font.italic = True

# Spacing after abstract
doc.add_paragraph()

def add_academic_paragraph(text):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 2.0  # Double spacing
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.first_line_indent = Inches(0.5) # Indented paragraphs
    run = p.add_run(text)
    run.font.size = Pt(12)
    return p

def add_heading_1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.bold = True
    run.font.size = Pt(12)
    return p

# 1. PENDAHULUAN (NO HEADING)
add_academic_paragraph(
    "Bidang pendidikan telah mengalami transformasi besar sebagai akibat dari kemajuan teknologi informasi, "
    "terutama dengan masuknya metode pembelajaran digital [1]. Learning Management System (LMS) adalah sebuah sistem "
    "digital yang dibuat untuk penyampaian materi pembelajaran secara jarak jauh dengan berbasis web dan mengelola aktivitas "
    "pembelajaran serta pencapaian hasilnya. Melalui pemanfaatan LMS, peserta memperoleh keleluasaan untuk menjangkau seluruh "
    "bahan ajar tanpa terikat oleh batasan waktu maupun lokasi. Pemanfaatan LMS dalam berbagai jenjang pendidikan terbukti "
    "meningkatkan efisiensi dan efektivitas interaksi antara pengajar dan pembelajar [3]."
)

add_academic_paragraph(
    "Bait Tahfizh Al-Qur'an Ridhallah (BTQR) merupakan lembaga pendidikan yang berfokus pada pengembangan Al-Qur'an "
    "dan bahasa Arab yang berlokasi di Kabupaten Aceh Besar, Provinsi Aceh. Sebagai salah satu lembaga terkemuka, BTQR memiliki "
    "sebuah program unggulan berupa bimbingan bahasa Arab untuk persiapan studi ke Timur Tengah yang diberi nama program "
    "Non-Reguler. Program bimbingan khusus ini sangat diminati oleh banyak peserta dari berbagai daerah di luar Aceh Besar. "
    "Namun demikian, kendala jarak lokasi yang sangat jauh menyulitkan mereka untuk rutin hadir mengikuti kegiatan pembelajaran "
    "secara tatap muka secara langsung [2]. Kendala ini mengakibatkan calon peserta terpaksa membatalkan niat mereka untuk bergabung. "
    "Oleh karena itu, pihak lembaga sangat membutuhkan solusi digital agar materi bimbingan tetap tersampaikan secara "
    "komprehensif tanpa mengharuskan kehadiran fisik peserta didik di lokasi."
)

add_academic_paragraph(
    "Untuk mengatasi permasalahan tersebut, kolaborasi riset dilakukan guna merancang dan membangun platform LMS berbasis web. "
    "Desain antarmuka pengguna (UI/UX) dirancang secara terpisah oleh Muhammad Faqih Sofyan yang berfokus pada analisis "
    "kegunaan (usability) dan keluhan pengguna. Sementara itu, penelitian ini berfokus pada implementasi arsitektur perangkat "
    "lunak dan fungsionalitas backend sistem. Pendekatan kolaboratif ini memastikan website yang dibangun memiliki kegunaan yang "
    "teruji serta fungsionalitas yang stabil."
)

add_academic_paragraph(
    "Sebagai solusi teknis, sistem dibangun berupa platform LMS berbasis web dengan mengintegrasikan Framework Laravel di sisi "
    "backend dan Vue.js di sisi frontend. Laravel dipilih karena memiliki kinerja tinggi, arsitektur Model-View-Controller "
    "(MVC) yang solid, serta fitur keamanan internal yang andal [11], [12]. Di sisi lain, Vue.js digunakan untuk menyediakan "
    "antarmuka pengguna yang dinamis, interaktif, dan responsif guna menyajikan pengalaman pengguna yang optimal [13], [14]. "
    "Integrasi kedua teknologi ini diharapkan mampu memberikan kemudahan aksesibilitas bagi peserta program bimbingan bahasa Arab "
    "di BTQR secara efisien dan aman. Tujuan utama penelitian ini adalah merancang dan mengembangkan Learning Management System "
    "(LMS) berbasis web menggunakan Laravel dan Vue.js guna mendukung proses bimbingan bahasa Arab di BTQR."
)

# 2. METODE
add_heading_1("METODE")

add_academic_paragraph(
    "Penelitian ini merupakan jenis penelitian dan pengembangan (Research and Development) yang difokuskan pada rekayasa "
    "perangkat lunak LMS. Metode pengembangan sistem yang diterapkan adalah model Waterfall yang bersifat linier dan sekuensial. "
    "Pemilihan model ini didasarkan pada kebutuhan sistem yang telah didefinisikan secara matang sejak awal fase proyek. "
    "Tahapan Waterfall dalam penelitian ini dijabarkan sebagai berikut:"
)

add_academic_paragraph(
    "Pertama, tahapan analisis kebutuhan (Requirements) yang melibatkan pengumpulan data kebutuhan fungsional dan "
    "non-fungsional sistem melalui observasi langsung dan wawancara dengan pengelola BTQR. Pada tahap ini pula, data spesifikasi "
    "dari penelitian kolaborator UI/UX diekstraksi untuk menjadi acuan fungsi antarmuka."
)

add_academic_paragraph(
    "Kedua, tahapan perancangan sistem (System Design) yang memetakan arsitektur database dan aliran fungsional aplikasi. "
    "Rancangan ini dimodelkan menggunakan Use Case Diagram untuk mendefinisikan batasan hak akses pengguna serta Entity "
    "Relationship Diagram (ERD) untuk menggambarkan struktur basis data relasional."
)

add_academic_paragraph(
    "Ketiga, tahapan implementasi (Implementation) yang berfokus pada penulisan kode program. Pembangunan logika bisnis "
    "dan RESTful API dilakukan menggunakan PHP dengan framework Laravel, sedangkan visualisasi antarmuka pengguna diimplementasikan "
    "menggunakan JavaScript dengan framework Vue.js."
)

add_academic_paragraph(
    "Keempat, tahapan pengujian (Testing) yang bertujuan memverifikasi seluruh komponen fungsional perangkat lunak. Pengujian "
    "dilakukan dengan menggunakan metode Black-box Testing untuk memastikan seluruh input menghasilkan output yang sesuai "
    "tanpa kegagalan sistem."
)

add_academic_paragraph(
    "Kelima, tahapan peluncuran (Deployment) yang berupa pengunggahan sistem ke server hosting (deployment) dan melakukan "
    "konfigurasi lingkungan produksi agar dapat diakses publik."
)

add_academic_paragraph(
    "Keenam, tahapan pemeliharaan (Maintenance) yang mencakup pemantauan berkala, perbaikan bug yang ditemukan pasca-rilis, "
    "dan optimasi performa berdasarkan umpan balik pengguna nyata."
)

add_academic_paragraph(
    "Subjek penelitian berpusat pada program Non-Reguler bimbingan bahasa Arab di Bait Tahfizh Al-Qur'an Ridhallah (BTQR), "
    "Aceh Besar, Aceh. Sumber data primer didapatkan dari pengelola lembaga (Admin), pengajar (Ustadz), dan siswa dari luar "
    "daerah. Sementara data sekunder bersumber dari dokumen elisitasi kebutuhan serta desain prototipe yang dihasilkan oleh "
    "Muhammad Faqih Sofyan."
)

# 3. HASIL
add_heading_1("HASIL")

add_academic_paragraph(
    "Pengembangan LMS bimbingan bahasa Arab di BTQR menghasilkan sebuah aplikasi berbasis web dinamis dengan tiga aktor utama: "
    "Admin, Pengajar, dan Siswa."
)

add_academic_paragraph(
    "Struktur database diimplementasikan secara relasional berdasarkan rancangan ERD yang mencakup tabel master `users` "
    "(dengan hak akses sebagai admin, pengajar, atau siswa), tabel `moduls` untuk menyimpan modul pembelajaran, tabel `materis` "
    "untuk menampung materi pembelajaran, tabel `try_outs` untuk menyimpan data ujian, serta tabel `try_out_hasil` untuk merekam "
    "skor dan kemajuan belajar siswa. Hubungan antartabel dioptimalkan dengan foreign key untuk menjamin integritas data di database MySQL."
)

add_academic_paragraph(
    "Visualisasi antarmuka sistem diwujudkan menggunakan Vue.js yang berkomunikasi secara asinkron dengan API Laravel. "
    "Representasi visual dari antarmuka ini digambarkan melalui teks berikut untuk memudahkan penyisipan gambar secara manual oleh penulis:"
)

# Text representation placeholders for images (single spaced, centered)
for i, desc in enumerate([
    "Gambar 1. Halaman Login Multi-User - Menampilkan form autentikasi terintegrasi dengan opsi login untuk admin, pengajar, dan siswa dengan desain modern minimalis menggunakan font Times New Roman.",
    "Gambar 2. Dashboard Admin - Memperlihatkan ringkasan data statistik jumlah siswa, pengajar, materi aktif, dan grafik perkembangan kelulusan try out.",
    "Gambar 3. Halaman Pengelolaan Modul oleh Pengajar - Menampilkan daftar materi bimbingan bahasa Arab (Nahwu, Sharaf, Mufradat) serta tombol pengelolaan untuk menambah, mengedit, atau menghapus modul.",
    "Gambar 4. Tampilan Ujian Try Out Siswa - Menampilkan lembar kerja interaktif bimbingan bahasa Arab untuk siswa yang sedang mengerjakan soal try out secara real-time lengkap dengan penunjuk waktu."
], 1):
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.paragraph_format.line_spacing = 1.15
    p_img.paragraph_format.space_after = Pt(12)
    run_img = p_img.add_run(f"[{desc}]")
    run_img.font.italic = True
    run_img.font.size = Pt(10)

add_academic_paragraph(
    "Verifikasi fungsionalitas sistem diuji menggunakan metode Black-box Testing. Hasil pengujian menunjukkan bahwa seluruh "
    "modul utama berjalan dengan sukses tanpa adanya kegagalan program. Ringkasan hasil pengujian disajikan pada Tabel 1."
)

# Table 1
p_tbl_title = doc.add_paragraph()
p_tbl_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_tbl_title.paragraph_format.space_after = Pt(6)
run_tbl_title = p_tbl_title.add_run("Tabel 1. Hasil Pengujian Fungsionalitas Menggunakan Black-Box Testing")
run_tbl_title.font.bold = True
run_tbl_title.font.size = Pt(10)

tbl_data = doc.add_table(rows=6, cols=5)
tbl_data.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl_data.autofit = True

headers = ["No", "Fitur/Modul", "Aktor Penguji", "Ekspektasi Output", "Status"]
rows_content = [
    ["1", "Manajemen Akun Pengguna", "Admin", "Akun baru berhasil dibuat, diubah, dan dihapus", "Berhasil"],
    ["2", "Kelola Modul & Materi", "Pengajar", "Modul ajar Nahwu & Sharaf berhasil diunggah", "Berhasil"],
    ["3", "Pelaksanaan Ujian Try Out", "Siswa", "Siswa dapat mengisi ujian dan nilai terhitung otomatis", "Berhasil"],
    ["4", "Belajar & Kuis Harian", "Siswa", "Konten materi tampil interaktif dan kuis tervalidasi", "Berhasil"],
    ["5", "Autentikasi & Otorisasi", "Semua Aktor", "Sistem membatasi hak akses sesuai role pengguna", "Berhasil"]
]

hdr_cells = tbl_data.rows[0].cells
for i, h_text in enumerate(headers):
    hdr_cells[i].text = h_text
    hdr_cells[i].paragraphs[0].runs[0].font.bold = True
    hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(10)
    set_cell_margins(hdr_cells[i], top=100, bottom=100, left=100, right=100)
    set_cell_borders(hdr_cells[i], color="999999", sz="4")

for r_idx, row in enumerate(rows_content, 1):
    row_cells = tbl_data.rows[r_idx].cells
    for c_idx, val in enumerate(row):
        row_cells[c_idx].text = val
        row_cells[c_idx].paragraphs[0].runs[0].font.size = Pt(10)
        set_cell_margins(row_cells[c_idx], top=100, bottom=100, left=100, right=100)
        set_cell_borders(row_cells[c_idx], color="DDDDDD", sz="2")

doc.add_paragraph()

# 4. PEMBAHASAN
add_heading_1("PEMBAHASAN")

add_academic_paragraph(
    "Berdasarkan hasil pengujian fungsionalitas, integrasi Laravel di sisi backend dan Vue.js di sisi frontend terbukti "
    "memberikan responsivitas aplikasi yang sangat baik dengan Separation of Concerns yang jelas [8]. Penggunaan Laravel "
    "memungkinkan pengelolaan database relasional seperti penanganan relasi rumit antara tabel `users`, `moduls`, dan `try_outs` "
    "menjadi sangat efisien berkat Eloquent ORM [11]. Di sisi lain, Vue.js berhasil menghadirkan interaksi yang mulus tanpa reload "
    "halaman penuh (SPA) yang secara signifikan meningkatkan kenyamanan belajar pengguna awam [13]."
)

add_academic_paragraph(
    "Model Waterfall memberikan panduan yang teratur sehingga setiap fase pengembangan terdokumentasi dengan baik dan "
    "meminimalisasi risiko penyimpangan fungsionalitas. Kolaborasi dengan perancang UI/UX (Muhammad Faqih Sofyan) terbukti "
    "mempercepat tahap implementasi karena pengembang tidak lagi melakukan trial-and-error pada tata letak visual, melainkan langsung "
    "berfokus pada efisiensi penulisan kode program dan ketahanan database."
)

add_academic_paragraph(
    "Jika dibandingkan dengan penelitian relevan sebelumnya, seperti LMS berbasis Laravel pada tingkat SMA [3] dan LMS pesantren "
    "tahfiz berbasis web [5], penelitian ini memberikan kontribusi spesifik pada integrasi pembelajaran bahasa Arab non-reguler dengan "
    "fokus pada materi Nahwu, Sharaf, dan evaluasi try out terintegrasi. Hal ini membuktikan bahwa pendekatan gabungan Laravel-Vue.js "
    "sangat relevan untuk diimplementasikan pada lembaga pendidikan keagamaan tradisional seperti BTQR."
)

# 5. KESIMPULAN DAN SARAN
add_heading_1("KESIMPULAN DAN SARAN")

add_academic_paragraph(
    "Berdasarkan proses pengembangan dan pengujian yang telah dilakukan, dapat disimpulkan bahwa Learning Management System (LMS) "
    "bimbingan bahasa Arab untuk Bait Tahfizh Al-Qur'an Ridhallah (BTQR) berhasil dibangun menggunakan Framework Laravel dan Vue.js. "
    "Sistem yang dikembangkan dengan metode Waterfall ini telah memenuhi seluruh spesifikasi kebutuhan fungsional, termasuk manajemen "
    "akun, pengelolaan modul materi, pelaksanaan ujian try out, serta evaluasi kemajuan belajar. Hasil pengujian dengan metode "
    "Black-box Testing menunjukkan bahwa seluruh fitur aplikasi berfungsi dengan sukses dan siap digunakan untuk menjembatani "
    "keterbatasan akses jarak jauh peserta luar kota."
)

add_academic_paragraph(
    "Sebagai saran untuk pengembangan selanjutnya, sistem ini direkomendasikan untuk menambahkan fitur gamifikasi guna "
    "meningkatkan motivasi belajar siswa [19], [20]. Selain itu, konversi platform ke aplikasi mobile berbasis hybrid (seperti "
    "Flutter atau React Native) dapat dipertimbangkan agar siswa dapat mengakses materi belajar secara lebih portabel melalui "
    "perangkat telepon pintar."
)

# 6. DAFTAR PUSTAKA
add_heading_1("DAFTAR PUSTAKA")

references_list = [
    "[1] Yauma, A., Fitri, I., & Ningsih, S. (2021). Learning Management System (LMS) pada E-Learning Menggunakan Metode Agile dan Waterfall berbasis Website. Jurnal Teknologi Informasi dan Komunikasi, 5(3), 323–328. https://doi.org/10.35870/jtik.v5i3.190",
    "[2] Fahmi, A. A., & Kholid, N. (2024). Manajemen Program Pembelajaran Bahasa Arab di Pesantren Darul Quran Aceh (DQA). Risalah: Jurnal Pendidikan dan Studi Islam, 10(4), 1519–1534.",
    "[3] Krisna, Ichwani, A., Anwar, N., & Sekti, B. A. (2024). Perancangan Learning Management System Pada SMA Islam Tambora. IKRA-ITH Informatika: Jurnal Komputer dan Informatika, 8(1), 201–212. https://doi.org/10.37817/ikraith-informatika.v8i1.3214",
    "[4] Ridha, M. R., Ayuni, S. K., & Shodiq, M. J. (2023). Pengembangan Media Learning Management System (LMS) Berbasis Kitāb Al-‘Arabiyah Li an-Nāsyi’īn. Al Mi’yar: Jurnal Ilmiah Pembelajaran Bahasa Arab dan Kebahasaaraban, 6(1), 1–12. https://doi.org/10.35931/am.v6i1.1842",
    "[5] Putra, A. D. A., & Rakhmadi, A. (2023). Sistem Monitoring Tahfidzul Qur’an Berbasis Web. Jurnal Teknik Informatika Universitas Muhammadiyah Surakarta, 4(2), 55-64.",
    "[6] Sulianta, F., Rumaisa, F., Puspitarani, Y., Violina, S., & Rosita, A. (2025). Abdimas Galuh a Learning Platform in the Digital Era. Jurnal Pengabdian Kepada Masyarakat, 7(2), 1720–1728.",
    "[7] Baehaqi, A., Basit, M. S., Indrajit, R. E., & Kurniawan, R. D. (2023). Front End Learning Management System Development Using the Nextjs Framework. Jurnal Teknik Informatika, 4(4), 899–911. https://doi.org/10.52436/1.jutif.2023.4.4.1273",
    "[8] Syafitri, A., Angraeni, A., & Wibowo, A. (2025). Sistem Informasi Manajemen Perpustakaan Digital Berbasis Web Menggunakan Framework Laravel. Jurnal Informatika dan Kesehatan, 2(2), 89–98. https://doi.org/10.35473/ikn.v2i2.3699",
    "[9] Wicaksono, G. W., Juliani, A. G., Wahyuni, E. D., Cholily, Y. M., Asrini, H. W., & Budiono. (2020). Analysis of Learning Management System Features based on Indonesian Higher Education National Standards using the Feature-Oriented Domain Analysis. Proceedings of 8th International Conference on Information and Communication Technology (ICoICT 2020). https://doi.org/10.1109/ICoICT49345.2020.9166459",
    "[10] Samihardjo, R., Murnawan, M., Amalia, E., & Pamungkas, A. C. (2023). Analysis of Web-Based E-Learning Management System Business Process to Increase Learning Effectiveness at SMAABC Bandung. Brilliance: Research of Artificial Intelligence, 3(2), 329–337. https://doi.org/10.47709/brilliance.v3i2.3274",
    "[11] Susilo, F. S., Hariyanto, E., & Utomo, R. B. (2024). Design and Development of a Learning Management System Information System using the Waterfall Method with the Laravel Framework. Journal of Systems Engineering and Information Technology, 5(4), 6084–6089.",
    "[12] Musyary, M., Kurniati, A., & Damarjati, C. (2023). Laravel Framework-Based Information System of the Department of Information Technology of Universitas Muhammadiyah Yogyakarta. Emerging Information Science and Technology, 4(2), 112-120. https://doi.org/10.18196/eist.v4i2.20736",
    "[13] Daar, G. F., Supartini, N. L., Sulasmini, N. M. A., Ekasani, K. A., Lestari, D., & Kesumayathi, I. A. G. (2023). Students’ Perception of the Use of Learning Management System in Learning English for Specific Purpose During the Pandemic: Evidence From Rural Area in Indonesia. Journal of Language Teaching and Research, 14(2), 403–409. https://doi.org/10.17507/jltr.1402.16",
    "[14] Anggraeni, O. S. I., Sugiarto, L., & Agustin, T. (2024). Studi Komparatif Performa Framework Javascript Modern dalam Pengembangan Aplikasi Web. Modem: Jurnal Informatika dan Sains Teknologi, 2(4), 162–177. https://doi.org/10.62951/modem.v2i4.239",
    "[15] Perdana, F. P., Supratman, E., & Saputra, R. R. (2024). Designing a Modern Web Interface with Vue.js and Tailwind for University Information System. Brilliance: Research of Artificial Intelligence, 4(2), 956–963. https://doi.org/10.47709/brilliance.v4i2.5409",
    "[16] Rahmani, A., Salistia, I. H., & Hizriani, N. (2023). Metode Pembelajaran Bahasa Arab Di Pondok Pesantren Nurul Hakim Kediri. Tatsqifiy: Jurnal Pendidikan Bahasa Arab, 4(1), 1–11. https://doi.org/10.30997/tjpba.v4i1.7100",
    "[17] Mujahidah, N., & Riyadhi, B. (2023). Model Pembelajaran Bahasa Arab di Pondok Pesantren. Jurnal Pendidikan Islam Al-Ilmi, 6(1), 22-31. https://doi.org/10.32529/al-ilmi.v6i1.2031",
    "[18] Amanta, A. N., & Dewi, M. U. (2026). Design and Development of STEKOM Learning Management System (SIAKAD) to Optimize User Experience. Journal of Applied Software Engineering, 5(1), 1–16.",
    "[19] Nugroho, R. P., Soepriyanto, Y., & Wedi, A. (2024). Development of Learning Management System with Gamification Approach for Project-Based Learning. JTP: Jurnal Teknologi Pendidikan, 26(3), 794–806. https://doi.org/10.21009/jtp.v26i3.40873",
    "[20] Hasanah, R. N. (2024). Rancang Bangun Learning Management System Berbasis Gamifikasi dengan Menggunakan Mechanics, Dynamics, Aesthetics (MDA) Framework. Jurnal Edutec, 6(2), 123-134.",
    "[21] Admaja, M. H., & Pramono, A. (2025). Pengembangan Learning Management System (LMS) Menggunakan Personal Extreme Programming (PxP) Pada SMP Negeri 1 Kedungadem. JATI: Jurnal Mahasiswa Teknik Informatika, 9(2), 3484–3492. https://doi.org/10.36040/jati.v9i2.13019"
]

for ref in references_list:
    p_ref = doc.add_paragraph()
    p_ref.paragraph_format.line_spacing = 1.15
    p_ref.paragraph_format.space_after = Pt(4)
    p_ref.paragraph_format.left_indent = Inches(0.25)
    p_ref.paragraph_format.first_line_indent = Inches(-0.25)
    run_ref = p_ref.add_run(ref)
    run_ref.font.size = Pt(10)

doc.save(r"c:\xampp\htdocs\BTQR\Paper\Artikel_LMS_BTQR.docx")
print("Successfully generated c:\\xampp\\htdocs\\BTQR\\Paper\\Artikel_LMS_BTQR.docx")
